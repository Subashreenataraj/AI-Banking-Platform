import hashlib
import re
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        document = DocxDocument(str(path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        tables = [
            " | ".join(cell.text for cell in row.cells)
            for table in document.tables
            for row in table.rows
        ]
        return "\n".join(paragraphs + tables)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"Unsupported document format: {suffix}")


def normalize_text(value: str) -> str:
    return re.sub(r"[ \t]+", " ", value.replace("\x00", "")).strip()


def content_hash(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()